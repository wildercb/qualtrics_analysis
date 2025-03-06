#!/usr/bin/env python
# coding: utf-8

# ### Imports

# In[3]:


from globals import survey_data, data_dir, SPECIAL_PHRASES

import os
import time
import re

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np

os.chdir(data_dir)



# ###########################################################
# ### Create quantitative analysis notebook 
# (Excel file that contains the count and percentage of all answers for all questions)
# 
# ###########################################################

# In[4]:


# Load the survey response data
file_path = survey_data
output_path = os.path.join(data_dir,'quantitative_analysis.xlsx')

df = pd.read_csv(file_path)


def protect_special_phrases(text: str, special_phrases) -> str:
    """
    Replace the commas in special phrases with a unique token, so we don't split them later.
    E.g. "other, please specify" -> "other<COMMA> please specify"
    """
    if not text:
        return text
    # Make a copy
    protected_text = text
    for phrase in special_phrases:
        # Replace the exact phrase "other, please specify" with "other<COMMA> please specify"
        # We do a case-sensitive replace here; adjust if you need case-insensitivity.
        safe_phrase = phrase.replace(",", "<COMMA>")  
        protected_text = protected_text.replace(phrase, safe_phrase)
    return protected_text

def restore_special_phrases(text: str) -> str:
    """
    Restore <COMMA> back to ',' in the text.
    """
    return text.replace("<COMMA>", ",")

# Function to sanitize sheet names for Excel
def sanitize_sheet_name(name):
    return re.sub(r'[\[\]*?:/\\]', '', name)[:31]

# Create a structured DataFrame to store the final results
summary_data = []

for column in df.columns:
    # Convert all values in this column to string (to handle NaN safely)
    responses = df[column].dropna().astype(str)
    
    # The first row is treated as the question text
    question_text = responses.iloc[0] if not responses.empty else ""
    
    # Exclude the first row from analysis
    responses = responses.iloc[1:]
    
    # 'respondent_count' is how many rows (respondents) actually answered this question
    respondent_count = len(responses)  # number of non-empty cells in this column (minus the question row)

    # Collect all answers in a list (splitting by comma for multiple choices)
    all_answers = []
    for response in responses:
        # 1) Protect special phrases in the response
        protected_response = protect_special_phrases(response, SPECIAL_PHRASES)
        
        # 2) Split by comma
        split_parts = [x.strip() for x in protected_response.split(',') if x.strip()]
        
        # 3) Restore the <COMMA> tokens in each part
        restored_parts = [restore_special_phrases(part) for part in split_parts]
        
        all_answers.extend(restored_parts)
    
    # Count how many times each distinct response was mentioned
    if len(all_answers) == 0:
        # No valid responses
        row_data = {
            "Question": f"{column}: {question_text}",
            "Respondent Count": 0
        }
        summary_data.append(row_data)
        continue

    response_counts = pd.Series(all_answers).value_counts()

    # If no one answered (respondent_count = 0), just record it and move on
    if respondent_count == 0:
        row_data = {
            "Question": f"{column}: {question_text}",
            "Respondent Count": 0
        }
        summary_data.append(row_data)
        continue

    # If there are more than 50 different types of responses, skip this question
    if len(response_counts) > 50:
        # Skip adding this row to the final output
        continue

    # Compute percentage based on the number of respondents
    response_percentages = (response_counts / respondent_count) * 100
    
    # Construct structured row
    row_data = {
        "Question": f"{column}: {question_text}",
        "Respondent Count": respondent_count
    }
    
    # Add each response's count & percentage to the row_data
    for i, (answer, count) in enumerate(response_counts.items(), start=1):
        row_data[f"Response {i}"] = answer
        row_data[f"Count {i}"] = count
        row_data[f"Percentage {i}"] = round(response_percentages[answer], 2)
    
    summary_data.append(row_data)

# Convert structured data into a DataFrame
final_df = pd.DataFrame(summary_data)

# Save analysis to an Excel file
final_df.to_excel(output_path, sheet_name='Survey Summary', index=False)

print(f"Analysis saved to {output_path}")


# #############################################################
# ### Create graphs for all questions and store in docx file
# ###############################################################

# In[5]:


import os
import re
import time
import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches

from selenium import webdriver
from selenium.webdriver.chrome.options import Options

# Define your data directory
xcel_path = 'quantitative_analysis.xlsx'

def save_fig_as_png_selenium(fig, filename, width=1200, height=700):
    """
    Save a Plotly figure as a PNG by:
      1) Writing an HTML file
      2) Opening it in a headless Chrome browser
      3) Taking a screenshot

    Note: Increased default width & height for more space.
    """
    html_path = filename.replace(".png", ".html")
    fig.write_html(html_path, include_plotlyjs="cdn", auto_open=False)

    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")

    driver = webdriver.Chrome(options=chrome_options)
    driver.set_window_size(width, height)

    driver.get("file://" + os.path.abspath(html_path))

    # Wait briefly for Plotly to renders
    driver.save_screenshot(filename)
    driver.quit()
    os.remove(html_path)

def sanitize_filename(name: str) -> str:
    """Remove or replace any invalid filename characters for Windows."""
    return re.sub(r'[\\/*?:"<>|]', '_', name)

if __name__ == "__main__":
    # Read your summary DataFrame
    df = pd.read_excel(xcel_path, sheet_name="Survey Summary")

    if not os.path.exists("figures"):
        os.makedirs("figures")

    doc = Document()

    for idx, row in df.iterrows():
        question_text = str(row["Question"])
        if pd.isna(question_text) or not question_text.strip():
            continue

        # Derive a short index for the PNG filename
        if ":" in question_text:
            question_index = question_text.split(":")[0].strip()
        else:
            question_index = f"Q{idx+1}"

        question_index = sanitize_filename(question_index)

        respondent_count = row.get("Respondent Count", 0)
        if pd.isna(respondent_count):
            respondent_count = 0

        # Collect response data from "Response X", "Count X", "Percentage X"
        responses = []
        i = 1
        while True:
            resp_col = f"Response {i}"
            count_col = f"Count {i}"
            pct_col = f"Percentage {i}"
            if resp_col not in df.columns:
                break
            val = row.get(resp_col)
            if pd.isna(val) or not str(val).strip():
                break
            count_val = row.get(count_col, 0)
            pct_val = row.get(pct_col, 0)
            if pd.isna(count_val):
                count_val = 0
            if pd.isna(pct_val):
                pct_val = 0
            responses.append({
                "Option": str(val),
                "Count": float(count_val),
                "Percentage": float(pct_val)
            })
            i += 1

        if not responses:
            continue

        data_df = pd.DataFrame(responses)

        # Prepare bar chart data
        x_vals = data_df["Option"]
        y_vals = data_df["Count"]
        # Text labels now just show the percentage value with '%' appended
        text_vals = [f"{pct}%" for pct in data_df["Percentage"]]

        # Create a color scale based on "Count"
        marker_dict = dict(
            color=y_vals,
            colorscale="Blues",
            showscale=True,
            colorbar=dict(title="Count")
        )

        fig = go.Figure()
        fig.add_trace(
            go.Bar(
                x=x_vals,
                y=y_vals,
                text=text_vals,
                textposition="outside",
                textfont=dict(
                    size=18,             # Larger font size for percentage text
                    color="black",
                    family="Arial Black" # Bold style font
                ),
                marker=marker_dict,
                cliponaxis=False  # Let text extend beyond the plot area
            )
        )

        # Compute a custom upper range for the y-axis to ensure no clipping
        max_y = y_vals.max() if not y_vals.empty else 0
        y_buffer = max_y * 0.15  # 15% buffer above highest bar

        # Update layout without a title for the graph image
        fig.update_layout(
            xaxis_title="Response Option",
            yaxis_title="Number of Mentions",
            template="plotly_white",
            xaxis=dict(
                tickangle=45,
                tickfont=dict(
                    family="Arial Black",  # Bold font for response options
                    size=12,
                    color="black"
                )
            ),
            margin=dict(l=80, r=80, t=80, b=80),
        )
        fig.update_yaxes(range=[0, max_y + y_buffer])

        # Save chart as PNG (via Selenium screenshot)
        image_path = os.path.join("figures", question_index + ".png")
        save_fig_as_png_selenium(fig, image_path, width=1200, height=700)

        # In the output docx, add a larger heading (level 1) above each graph.
        doc.add_heading(question_text, level=2)
        doc.add_picture(image_path, width=Inches(6))

        doc.add_paragraph(f"Number of respondents: {int(respondent_count)}")
        doc.add_paragraph("Breakdown of responses:")
        for _, r in data_df.iterrows():
            response_line = f"{r['Option']}: {int(r['Count'])} ({r['Percentage']}%)"
            doc.add_paragraph(response_line)
        doc.add_paragraph("")

    doc.save("quantitative_analysis.docx")
    print("Report generated")


# ######################################
# ### Create Demographics Table 
# ######################################

# In[ ]:


######################
# Table and Stats
######################

import seaborn as sns
from tabulate import tabulate
from matplotlib.table import Table
import matplotlib.font_manager as fm
from IPython.display import display, Image

input_file = "survey_data.csv"
plot_title = "Demographics All"
# Try different encodings to read the data file
encodings = ['utf-8', 'ISO-8859-1', 'windows-1252']
for encoding in encodings:
    try:
        df = pd.read_csv(input_file, encoding=encoding)

        track_a_columns = list(range(38, 65))  # AM is the 39th column (index 38), BM is the 65th column (index 64)

        # Function to check if a row belongs to Track A
        def is_track_a(row):
            return row[track_a_columns].notna().any()

        df = df.drop([0]).reset_index(drop=True)
        # Count Track A and Track B participants
        track_a_count = df[df.apply(is_track_a, axis=1)].shape[0]
        track_b_count = df.shape[0] - track_a_count
        total_count = df.shape[0]

        # Calculate percentages
        track_a_percent = (track_a_count / total_count) * 100
        track_b_percent = (track_b_count / total_count) * 100

        # Print the results
        print(f"Number of Track A participants: {track_a_count} ({track_a_percent:.2f}%)")
        print(f"Number of Track B participants: {track_b_count} ({track_b_percent:.2f}%)")
        print(f"Total number of participants: {total_count}")
        print(f"Successfully read the file with {encoding} encoding.")
        break
    except UnicodeDecodeError:
        print(f"Failed to read with {encoding} encoding. Trying next...")
else:
    print("Failed to read the file with any of the specified encodings.")
    raise

# Step 1: Clean and Rename Columns
demographic_columns = ['P1', 'P2', 'P3', 'P4', 'P5', 'P6', 'P7', 'P8', 'P9', 'P10', 'P11', 'P12']
demographics_df = df[demographic_columns]

# Remove the first two rows (metadata) and reset the index
demographics_cleaned_df = demographics_df#.drop([0]).reset_index(drop=True)

# Rename the columns to reflect actual demographic categories
demographics_cleaned_df.columns = [
    'Age', 'Gender', 'Education_Level', 'Field_of_Study', 'Company_Size',
    'Company_Type', 'Location', 'Role', 'Experience_Years_At_Company', 
    'Experience_Years_In_Development', 'Develop_AI_Products', 'AI_User'
]

# Function to group locations
def group_location(location):
    if pd.isna(location):
        return 'Empty/NaN'
    elif location == 'North America':
        return 'North America'
    elif location == 'EU/UK/EEA':
        return 'EU/UK/EEA'
    elif location == 'Central/South America':
        return 'Central/South America'
    else:
        return 'World'

# Apply the grouping function to the Location column
demographics_cleaned_df['Grouped_Location'] = demographics_cleaned_df['Location'].apply(group_location)

# Handle "None" and NaN values in 'Experience_Years_In_Development'
demographics_cleaned_df['Experience_Years_In_Development'] = (
    demographics_cleaned_df['Experience_Years_In_Development']
    .replace([np.nan, 'nan'], 'None')  # Treat NaNs as 'None'
    .astype(str)  # Ensure all values are strings for consistent comparison
    .str.strip()  # Remove any leading or trailing whitespace
)

# Function to create demographic table as an image
def create_demographic_table_image(data):
    # Calculate percentages for each category
    gender_counts = data['Gender'].value_counts(normalize=True) * 100
    age_counts = data['Age'].value_counts(normalize=True) * 100
    education_counts = data['Education_Level'].value_counts(normalize=True) * 100
    
    bachelors_count = sum(count for level, count in education_counts.items() if "bachelor" in str(level).lower())
    masters_count = sum(count for level, count in education_counts.items() if "master" in str(level).lower())
    phd_count = education_counts.get('Ph.D.', 0)
    high_school_count = education_counts.get('High School Degree', 0)
    other_education = 100 - (bachelors_count + masters_count + phd_count + high_school_count)

    field_counts = data['Field_of_Study'].value_counts(normalize=True) * 100
    company_size_counts = data['Company_Size'].value_counts(normalize=True) * 100
    company_type_counts = data['Company_Type'].value_counts(normalize=True) * 100
    location_counts = data['Grouped_Location'].value_counts(normalize=True) * 100
    role_counts = data['Role'].value_counts(normalize=True) * 100
    exp_dev_counts = data['Experience_Years_In_Development'].value_counts(normalize=True) * 100
    exp_company_counts = data['Experience_Years_At_Company'].value_counts(normalize=True) * 100

    # Group roles
    ai_manager = role_counts.get('Administrative role (CEO, Chief Technical Officer, Chief Operating Officer, Chief Information Officer)', 0) + role_counts.get('AI Manager', 0)
    requirements_analyst = role_counts.get('Requirements Analyst or Engineer', 0) + role_counts.get('Scrum Master, Product Manager, or Project Manager', 0)
    ai_developer = role_counts.get('AI Engineer or Developer', 0) + role_counts.get('(Software) Developer, Designer, or Architect', 0) + role_counts.get('Data Scientist or Data Analyst', 0)
    security_privacy = role_counts.get('Information Security Analyst or Engineer', 0) + role_counts.get('Information Privacy Analyst or Engineer', 0)
    other_roles = 100 - (ai_manager + requirements_analyst + ai_developer + security_privacy)
    total_respondents = len(data)
    role_groups = {
        'AI Manager': ['Administrative role (CEO, Chief Technical Officer, Chief Operating Officer, Chief Information Officer)', 'AI Manager'],
        'Requirements Analyst': ['Requirements Analyst or Engineer', 'Scrum Master, Product Manager, or Project Manager'],
        'AI Developer': ['AI Engineer or Developer', '(Software) Developer, Designer, or Architect', 'Data Scientist or Data Analyst'],
        'Security / Privacy': ['Information Security Analyst or Engineer', 'Information Privacy Analyst or Engineer'],
        'AR' : ['AI Researcher','AI Ethicist']
    }

    grouped_roles = {}
    for group, roles in role_groups.items():
        count = sum(data['Role'].value_counts().get(role, 0) for role in roles)
        percentage = (count / total_respondents) * 100
        grouped_roles[group] = (count, percentage)

    # Calculate 'Other' category
    other_count = total_respondents - sum(count for count, _ in grouped_roles.values())
    other_percentage = (other_count / total_respondents) * 100
    grouped_roles['Other'] = (other_count, other_percentage)

    print(f"\nGrouped Roles (Total Respondents: {total_respondents}):")
    for role, (count, percentage) in grouped_roles.items():
        print(f"{role}: {count} respondents ({percentage:.1f}%)")

    # Calculate percentages for the entire dataset
    role_counts = data['Role'].value_counts()
    total_roles = role_counts.sum()

    ai_manager = sum(role_counts.get(role, 0) for role in role_groups['AI Manager']) / total_roles * 100
    requirements_analyst = sum(role_counts.get(role, 0) for role in role_groups['Requirements Analyst']) / total_roles * 100
    ai_developer = sum(role_counts.get(role, 0) for role in role_groups['AI Developer']) / total_roles * 100
    security_privacy = sum(role_counts.get(role, 0) for role in role_groups['Security / Privacy']) / total_roles * 100
    other_roles = 100 - (ai_manager + requirements_analyst + ai_developer + security_privacy)

    print("\nPercentages for the entire dataset:")
    print(f"AI Manager: {ai_manager:.1f}%")
    print(f"Requirements Analyst: {requirements_analyst:.1f}%")
    print(f"AI Developer: {ai_developer:.1f}%")
    print(f"Security / Privacy: {security_privacy:.1f}%")
    print(f"Other: {other_roles:.1f}%")
    # Create the table data
    table_data = [
        ["Gender", f"Female ({gender_counts.get('Female', 0):.1f}%)", 
                   f"Male ({gender_counts.get('Male', 0):.1f}%)", 
                   f"Non-Binary ({gender_counts.get('Non-binary / Third gender', 0):.1f}%)", 
                   f"Other ({gender_counts.get('Other', 0):.1f}%)",
                   f"PnS ({gender_counts.get('Prefer not to say', 0):.1f}%)"],
        ["Age", f"18-25 ({age_counts.get('18-25', 0):.1f}%)", 
               f"26-35 ({age_counts.get('26-35', 0):.1f}%)", 
               f"36-45 ({age_counts.get('36-45', 0):.1f}%)", 
               f"46+ ({age_counts.get('46+', 0):.1f}%)",
               f"PnS ({age_counts.get('Prefer not to say', 0):.1f}%)"],
        ["Education", f"High School ({high_school_count:.1f}%)", 
                      f"Bachelor's ({bachelors_count:.1f}%)", 
                      f"Master's ({masters_count:.1f}%)",
                      f"Ph.D. ({phd_count:.1f}%)",
                      f"Other ({other_education:.1f}%)"],
        ["Degree Field", f"CS/ECE ({field_counts.get('Computer Science', 0) + field_counts.get('Electrical and Computer Engineering', 0):.1f}%)",
                  f"SWE/DS ({field_counts.get('Software Engineering', 0) + field_counts.get('Data Science', 0):.1f}%)",  
                  f"IT/Info Sec ({field_counts.get('Information Science (e.g., IT or MIS)', 0) + field_counts.get('Information Security/Privacy Engineer', 0):.1f}%)", 
                  f"Business ({field_counts.get('Business', 0):.1f}%)", 
                  f"Other ({field_counts.get('Other, please specify', 0):.1f}%)"],
        ["Company Size", f"1-5 ({company_size_counts.get('1-5 Employees', 0):.1f}%)", 
                         f"6-20 ({company_size_counts.get('6-20 Employees', 0):.1f}%)", 
                         f"21-50 ({company_size_counts.get('21-50 Employees', 0):.1f}%)",
                         f"51-100 ({company_size_counts.get('51-100 Employees', 0):.1f}%)",
                         f"100+ ({company_size_counts.get('101+ Employees', 0):.1f}%)"],
        ["Company Type", f"Multi-national ({company_type_counts.get('Multi-national Corporate', 0):.1f}%)", 
                         f"Startup/Small ({company_type_counts.get('Startup/Small Business', 0):.1f}%)", 
                         f"Academic/Research ({company_type_counts.get('Academic Institution/Research Center', 0):.1f}%)", 
                         f"Government ({company_type_counts.get('Government', 0):.1f}%)",
                         f"Other ({company_type_counts.get('Individual', 0) + field_counts.get('Other, please specify', 0):.1f}%)"],
        ["Location", f"N. America ({location_counts.get('North America', 0):.1f}%)", 
                     f"EU/UK/EEA ({location_counts.get('EU/UK/EEA', 0):.1f}%)", 
                     f"C/S America ({location_counts.get('Central/South America', 0):.1f}%)", 
                     f"World ({location_counts.get('World', 0):.1f}%)"],
        ["Role", f"AI Manager ({ai_manager:.1f}%)", 
                 f"Requirements Analyst ({requirements_analyst:.1f}%)", 
                 f"AI Developer ({ai_developer:.1f}%)", 
                 f"Security / Privacy ({security_privacy:.1f}%)", 
                 f"Other ({other_roles:.1f}%)"],
        ["Dev. Exp.", f"None ({exp_dev_counts.get('None', 0):.1f}%)", 
                    f"1-2 yrs ({exp_dev_counts.get('1-2 Years', 0):.1f}%)", 
                    f"2-5 yrs ({exp_dev_counts.get('2-5 Years', 0):.1f}%)", 
                    f"5-10 yrs ({exp_dev_counts.get('5-10 Years', 0):.1f}%)", 
                    f"10+ yrs ({exp_dev_counts.get('10+ Years', 0):.1f}%)"],
        ["Company Exp.", f"1-2 yrs ({exp_company_counts.get('1-2 Years', 0):.1f}%)", 
                         f"2-5 yrs ({exp_company_counts.get('2-5 Years', 0):.1f}%)", 
                         f"5-10 yrs ({exp_company_counts.get('5-10 Years', 0):.1f}%)", 
                         f"10+ yrs ({exp_company_counts.get('10+ Years', 0):.1f}%)"]
    ]
    # Ensure all rows have the same number of columns
    max_cols = max(len(row) for row in table_data)
    table_data = [row + [''] * (max_cols - len(row)) for row in table_data]
    
    # Create figure and axis for the table image
    fig, ax = plt.subplots(figsize=(12, 2))
    ax.axis('off')
    
    # Create and style table
    table = Table(ax, bbox=[0, 0, 1, 1])
    for (row, col), cell in np.ndenumerate(table_data):
        color = 'white'
        cell_obj = table.add_cell(row, col, width=1/max_cols, height=1/len(table_data), text=cell, 
                                  loc='center', facecolor=color)
        cell_obj.set_text_props(fontproperties=fm.FontProperties(family='Times New Roman', size=6))
        if col == 0:
            cell_obj.set_text_props(fontproperties=fm.FontProperties(family='Times New Roman', size=6, weight='bold'))
    ax.add_table(table)
    
    plt.title(plot_title, 
              fontproperties=fm.FontProperties(family='Times New Roman', size=10, weight='bold'),
              y=1.05)
    
    # Save and display the image
    plt.savefig('demographic_table.png', dpi=300, bbox_inches='tight')
    plt.close()

    # Display the table as text in addition to the PNG image
    print(tabulate(table_data, headers="firstrow", tablefmt="pipe"))

    return 'demographic_table.png'

# Create and display demographic table
image_path = create_demographic_table_image(demographics_cleaned_df)
display(Image(filename=image_path))

# Function to print demographic summaries in text format
def display_demographic_summary(data):
    total_respondents = len(data)
    for column in data.columns:
        print(f"\n--- {column.replace('_', ' ')} --- (Total Respondents: {total_respondents})")
        value_counts = data[column].value_counts(dropna=False)
        percentages = (value_counts / total_respondents) * 100

        for value, count in value_counts.items():
            if pd.isna(value):
                print(f"Empty/NaN: {count} respondents ({percentages[value]:.2f}%)")
            else:
                print(f"{value}: {count} respondents ({percentages[value]:.2f}%)")

        empty_count = data[column].isna().sum()
        empty_percentage = (empty_count / total_respondents) * 100
        print(f"Total empty responses: {empty_count} ({empty_percentage:.2f}%)")

def create_role_breakdown_table(data):
    role_groups = {
        'AM: AI Manager, Administrative Role': ['Administrative role (CEO, Chief Technical Officer, Chief Operating Officer, Chief Information Officer)', 'AI Manager'],
        'AD: AI/Software Designer, Architect, Developer, Data Scientist': ['AI Engineer or Developer', '(Software) Developer, Designer, or Architect', 'Data Scientist or Data Analyst'],
        'RA: Requirement Analyst, Product Manager': ['Scrum Master, Product Manager, or Project Manager','Requirements Analyst or Engineer'],
        'Sec: Information Security/Privacy Analyst or Engineer': ['Information Security Analyst or Engineer', 'Information Privacy Analyst or Engineer'],
        'AR: AI Researcher': ['AI Researcher']
    }

    grouped_roles = {}
    for group, roles in role_groups.items():
        count = sum(data['Role'].value_counts().get(role, 0) for role in roles)
        grouped_roles[group] = count

    # Calculate the 'Other' category
    total_count = 414  # we should probably do this better 
    other_count = total_count - sum(grouped_roles.values())
    grouped_roles['Other'] = other_count

    # Create the table data
    table_data = [
        ['Role', 'Count'],
        *[[role, str(count)] for role, count in grouped_roles.items()],
        ['Total', str(total_count)]
    ]

    # Create figure and axis for the table image
    fig, ax = plt.subplots(figsize=(4, 1.2))
    ax.axis('off')
    
    # Create and style table
    table = Table(ax, bbox=[0, 0, 1, 1])
    n_rows = len(table_data)
    width_ratios = [0.85, 0.15]  # 85% for Role, 15% for Count

    for (row, col), cell in np.ndenumerate(table_data):
        color = 'white'
        cell_text = cell if col == 0 else f'{cell:>5}'
        cell_obj = table.add_cell(row, col, width_ratios[col], 1/n_rows, text=cell_text, 
                                  loc='left' if col == 0 else 'right', facecolor=color)
        cell_obj.set_text_props(fontproperties=fm.FontProperties(family='Times New Roman', size=6))
        if col == 0:
            cell_obj.set_text_props(fontproperties=fm.FontProperties(family='Times New Roman', size=6, weight='bold'))
            cell_obj.PAD = 0.02
    ax.add_table(table)
    
    plt.title("Table 1: Breakdown of Participants' Roles", 
              fontproperties=fm.FontProperties(family='Times New Roman', size=8, weight='bold'),
              y=1.05)
    
    # Save and display the image
    plt.savefig('role_breakdown_table.png', dpi=300, bbox_inches='tight')
    plt.close()

    print("Table image saved as 'role_breakdown_table.png'")

    # Display the table as text in addition to the PNG image
    print(tabulate(table_data, headers="firstrow", tablefmt="pipe"))

# Call the function
create_role_breakdown_table(demographics_cleaned_df)

# Display the image
display(Image('role_breakdown_table.png'))
# Display demographic summary
display_demographic_summary(demographics_cleaned_df)

# Filter the dataframe for participants with Business degree
business_df = demographics_cleaned_df[demographics_cleaned_df['Field_of_Study'] == 'Business']

# Calculate the percentage of Business degree holders
total_participants = len(demographics_cleaned_df)
business_participants = len(business_df)
business_percentage = (business_participants / total_participants) * 100

print(f"Participants with Business degree: {business_participants} out of {total_participants} ({business_percentage:.2f}%)")

# Analyze role distribution for Business degree holders
role_distribution = business_df['Role'].value_counts()
role_percentages = (role_distribution / business_participants) * 100

print("\nRole distribution for Business degree holders:")
for role, count in role_distribution.items():
    percentage = role_percentages[role]
    print(f"{role}: {count} ({percentage:.2f}%)")


# New cell for "Other, please specify" Field of Study analysis

# Filter the dataframe for participants who selected "Other, please specify" for Field of Study
other_field_df = demographics_cleaned_df[demographics_cleaned_df['Field_of_Study'] == 'Other, please specify']

# Calculate the percentage of "Other, please specify" participants
total_participants = len(demographics_cleaned_df)
other_field_participants = len(other_field_df)
other_field_percentage = (other_field_participants / total_participants) * 100

print(f"Participants with 'Other, please specify' Field of Study: {other_field_participants} out of {total_participants} ({other_field_percentage:.2f}%)")

# Analyze role distribution for "Other, please specify" Field of Study participants
role_distribution = other_field_df['Role'].value_counts()
role_percentages = (role_distribution / other_field_participants) * 100

print("\nRole distribution for 'Other, please specify' Field of Study participants:")
for role, count in role_distribution.items():
    percentage = role_percentages[role]
    print(f"{role}: {count} ({percentage:.2f}%)")

